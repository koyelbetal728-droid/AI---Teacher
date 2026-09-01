# document_loader.py
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentLoader:
    """
    Loads text from common educational document formats.

    Supported formats:
    - PDF
    - DOCX
    - TXT
    - MD

    The loader only extracts text. Cleaning, chunking,
    embeddings, and vector storage are handled by the
    other RAG modules.
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
        ".md",
    }

    # ---------------------------------------------------------
    # Load document
    # ---------------------------------------------------------

    def load(
        self,
        file_path: str | Path,
    ) -> dict[str, Any]:
        """
        Load a document and return its extracted text
        together with basic metadata.
        """

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Document not found: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {path}"
            )

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported document type: {extension}. "
                f"Supported types: "
                f"{', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        if extension == ".pdf":
            text, metadata = self._load_pdf(path)

        elif extension == ".docx":
            text, metadata = self._load_docx(path)

        elif extension in {".txt", ".md"}:
            text, metadata = self._load_text(path)

        else:
            raise ValueError(
                f"Unsupported document type: {extension}"
            )

        return {
            "text": text,
            "metadata": {
                "filename": path.name,
                "file_path": str(path),
                "extension": extension,
                **metadata,
            },
        }

    # ---------------------------------------------------------
    # PDF
    # ---------------------------------------------------------

    def _load_pdf(
        self,
        path: Path,
    ) -> tuple[str, dict[str, Any]]:
        """
        Extract text from a PDF.
        """

        reader = PdfReader(str(path))

        pages = []
        page_count = len(reader.pages)

        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""

            page_text = page_text.strip()

            if page_text:
                pages.append(
                    f"[Page {page_number}]\n"
                    f"{page_text}"
                )

        text = "\n\n".join(pages)

        return text, {
            "page_count": page_count,
            "extracted_pages": len(pages),
        }

    # ---------------------------------------------------------
    # DOCX
    # ---------------------------------------------------------

    def _load_docx(
        self,
        path: Path,
    ) -> tuple[str, dict[str, Any]]:
        """
        Extract paragraphs and table contents from DOCX.
        """

        document = DocxDocument(
            str(path)
        )

        sections = []

        # Paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                sections.append(text)

        # Tables
        for table in document.tables:
            rows = []

            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                if any(cells):
                    rows.append(
                        " | ".join(cells)
                    )

            if rows:
                sections.append(
                    "\n".join(rows)
                )

        return "\n\n".join(sections), {
            "paragraph_count": len(
                document.paragraphs
            ),
            "table_count": len(
                document.tables
            ),
        }

    # ---------------------------------------------------------
    # TXT / Markdown
    # ---------------------------------------------------------

    def _load_text(
        self,
        path: Path,
    ) -> tuple[str, dict[str, Any]]:
        """
        Read plain-text or Markdown files.
        """

        encodings = [
            "utf-8",
            "utf-8-sig",
            "latin-1",
        ]

        text = None
        used_encoding = None

        for encoding in encodings:
            try:
                text = path.read_text(
                    encoding=encoding
                )
                used_encoding = encoding
                break

            except UnicodeDecodeError:
                continue

        if text is None:
            raise ValueError(
                "Unable to decode text file."
            )

        return text.strip(), {
            "encoding": used_encoding,
        }

    # ---------------------------------------------------------
    # Load multiple documents
    # ---------------------------------------------------------

    def load_many(
        self,
        file_paths: list[str | Path],
    ) -> list[dict[str, Any]]:
        """
        Load multiple documents.
        """

        documents = []

        for file_path in file_paths:
            documents.append(
                self.load(file_path)
            )

        return documents

    # ---------------------------------------------------------
    # Get metadata
    # ---------------------------------------------------------

    def get_file_metadata(
        self,
        file_path: str | Path,
    ) -> dict[str, Any]:
        """
        Return basic metadata without extracting content.
        """

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Document not found: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {path}"
            )

        extension = path.suffix.lower()

        return {
            "filename": path.name,
            "file_path": str(path),
            "extension": extension,
            "size_bytes": path.stat().st_size,
            "supported": (
                extension
                in self.SUPPORTED_EXTENSIONS
            ),
        }


# -------------------------------------------------------------
# Default loader
# -------------------------------------------------------------

document_loader = DocumentLoader()